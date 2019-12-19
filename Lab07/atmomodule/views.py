from django.shortcuts import render
from django.http import FileResponse
from .models import Sensor, SensorCategory, AtmoSnapshot
from django.views.generic import CreateView
from dicttoxml import dicttoxml
from xml.dom.minidom import parseString
from django.forms.models import model_to_dict
from docx import Document
from docx.shared import Cm
import datetime


def index(request):
    return render(request, "index.html")

def sensor(request, id):
    sensor = Sensor.objects.filter(id=id).values()
    category = SensorCategory.objects.filter(id=list(sensor)[0]['category_id']).values()
    return render(request, "sensor.html", context={"sensor": list(sensor)[0], "category": list(category)[0]})

def sensors(request):
    sensors = Sensor.objects.all()
    return render(request, "sensors.html", context={"sensors": list(sensors)})

def exports(request):
    return render(request, "exports.html")

def export_xml(request):
    try:
        exporter = Exporter()
        exporter.export_to_xml(get_data())
        return render(request, "success.html", context={"data": 'xml'})
    except Exception:
        return render(request, "fail.html", context={"data": 'xml'})

def export_docx(request):
    try:
        exporter = Exporter()
        exporter.export_to_docx(get_data())
        return render(request, "success.html", context={"data": 'docx'})
    except Exception:
        return render(request, "fail.html", context={"data": 'docx'})

def success(request, data):
    return render(request, "success.html", context={"data": data})

def get_xml(request):
    try:
        return render(request, 'formated_data\\data.xml', content_type='text/xml')
    except Exception:
        return render(request, "fail.html", context={"data": 'xml_get'})

def get_docx(request):
    try:
        return FileResponse(open('templates\\formated_data\\data.docx', 'rb'))
    except Exception:
        return render(request, "fail.html", context={"data": 'docx_get'})

class SensorCategoryCreateView(CreateView):
    template_name = 'creations/sensor_category_create.html'
    success_url="/success/sensor_category/"
    model = SensorCategory
    fields = ('name',)

class SensorCreateView(CreateView):
    template_name = 'creations/sensor_create.html'
    success_url="/success/sensor/"
    model = Sensor
    fields = ('lon', 'lat', 'create_date', 'name', 'model', 'category',)

class AtmoSnapshotCreateView(CreateView):
    template_name = 'creations/snapshot_create.html'
    success_url="/success/snapshot/"
    model = AtmoSnapshot
    fields = ('temperature', 'pressure', 'co2_level', 'sensor',)




def get_data():
    data_model = ModelData()
    sensor_categories = SensorCategory.objects.all()
    sensors = Sensor.objects.all()
    snapshots = AtmoSnapshot.objects.all()
    data_model.add_data(sensor_categories, "Sensor Categories")
    data_model.add_data(sensors, "Sensors")
    data_model.add_data(snapshots, "Snapshots")
    data_model.map("Sensors", "Sensor Categories", "category", "name")
    data_model.map("Snapshots", "Sensors", "sensor", "name")
    return data_model.get_data()

class Exporter(object):

    def export_to_xml(self, data):
        xml = dicttoxml(data, custom_root='Data', attr_type=False)
        dom = parseString(xml)
        file = open('templates\\formated_data\\data.xml', 'w')
        file.write(dom.toprettyxml())
        file.close()

    def export_to_docx(self, data):
        document = Document()
        document.add_heading('Data of AtmoStation', 0)
        document.add_paragraph(str(datetime.datetime.now()))
        for k, v in data.items():
            document.add_heading(k, level=1)
            for v1 in v:
                document.add_heading("Item", level=2).paragraph_format.left_indent = Cm(0.5)
                for v2 in v1:
                    document.add_paragraph(str(v2) + ": " + str(v1[v2]), style='List Bullet').paragraph_format.left_indent = Cm(1.5)

        document.save('templates\\formated_data\\data.docx')

class ModelData(object):
    def __init__(self):
        self.data = {}

    def add_data(self, data, name):
        self.__convert_to_dict(data, name)

    def set_data(self, data):
        self.data = data

    def get_data(self):
        return self.data

    def map(self, name_data1, name_data2, filed_cmp, field_fill):
        for d in self.data[name_data1]:
            d[filed_cmp] = next(d2 for d2 in self.data[name_data2] if d2["id"]==d[filed_cmp])[field_fill]

    def __convert_to_dict(self, data, name):
        cur_data = []
        for d in data:
            cur_data.append(model_to_dict(d))
        self.data[name] = cur_data